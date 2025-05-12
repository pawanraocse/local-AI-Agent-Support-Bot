import os
from typing import List, Dict
import logging
import json
from config import OLLAMA_HOST, MODEL_NAME
import PyPDF2
import docx
import numpy as np
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import DocArrayInMemorySearch
from langchain.text_splitter import RecursiveCharacterTextSplitter
from datetime import datetime
import traceback
import warnings

# Suppress Pydantic deprecation warning
warnings.filterwarnings("ignore", message="`pydantic.error_wrappers:ValidationError` has been moved to `pydantic:ValidationError`")

# Configure logging with structured format
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, documents_dir: str = "/app/documents"):
        logger.info(f"Initializing DocumentProcessor with directory: {documents_dir}")
        self.documents_dir = documents_dir
        self.embeddings = OllamaEmbeddings(model=MODEL_NAME, base_url=OLLAMA_HOST)
        self.vectorstore = None
        self.chunk_size = 1000
        self.chunk_overlap = 200
        self.processed_files = {}
        self.metadata_file = os.path.join(documents_dir, "processed_files.json")
        self._load_processed_files()
        self.scan_and_process_documents()

    def _load_processed_files(self):
        """Load the record of processed files."""
        try:
            if os.path.exists(self.metadata_file):
                with open(self.metadata_file, 'r') as f:
                    self.processed_files = json.load(f)
                logger.info(f"Loaded metadata for {len(self.processed_files)} processed files from {self.metadata_file}")
            else:
                logger.info(f"No existing metadata file found at {self.metadata_file}")
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in metadata file: {str(e)}")
            self.processed_files = {}
        except Exception as e:
            logger.error(f"Error loading processed files metadata: {str(e)}\n{traceback.format_exc()}")
            self.processed_files = {}

    def _save_processed_files(self):
        """Save the record of processed files."""
        try:
            with open(self.metadata_file, 'w') as f:
                json.dump(self.processed_files, f, indent=2)
            logger.info(f"Saved metadata for {len(self.processed_files)} files to {self.metadata_file}")
        except Exception as e:
            logger.error(f"Error saving processed files metadata: {str(e)}\n{traceback.format_exc()}")

    def _get_file_metadata(self, file_path: str) -> Dict:
        """Get metadata for a file (modification time and size)."""
        try:
            stat = os.stat(file_path)
            metadata = {
                "modified_time": stat.st_mtime,
                "size": stat.st_size,
                "last_processed": datetime.now().isoformat()
            }
            logger.debug(f"Retrieved metadata for {file_path}: {metadata}")
            return metadata
        except Exception as e:
            logger.error(f"Error getting metadata for {file_path}: {str(e)}")
            return {}

    def _is_file_modified(self, file_path: str) -> bool:
        """Check if a file has been modified since last processing."""
        try:
            if file_path not in self.processed_files:
                logger.info(f"New file detected: {file_path}")
                return True
            
            current_metadata = self._get_file_metadata(file_path)
            stored_metadata = self.processed_files[file_path]
            
            is_modified = (current_metadata["modified_time"] != stored_metadata["modified_time"] or
                         current_metadata["size"] != stored_metadata["size"])
            
            if is_modified:
                logger.info(f"File modified since last processing: {file_path}")
                logger.debug(f"Current metadata: {current_metadata}")
                logger.debug(f"Stored metadata: {stored_metadata}")
            
            return is_modified
        except Exception as e:
            logger.error(f"Error checking file modification status for {file_path}: {str(e)}")
            return True  # Process the file if there's an error checking its status

    def scan_and_process_documents(self):
        """Scan documents directory and process new or modified files."""
        logger.info(f"Starting document scan in directory: {self.documents_dir}")
        all_chunks = []
        files_processed = False
        processed_count = 0
        skipped_count = 0
        error_count = 0

        try:
            # Walk through the documents directory
            for root, _, files in os.walk(self.documents_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    
                    # Skip the metadata file
                    if file_path == self.metadata_file:
                        logger.debug(f"Skipping metadata file: {file_path}")
                        continue

                    try:
                        if self._is_file_modified(file_path):
                            logger.info(f"Processing file: {file_path}")
                            chunks = self.load_document(file_path)
                            if chunks:
                                all_chunks.extend(chunks)
                                self.processed_files[file_path] = self._get_file_metadata(file_path)
                                files_processed = True
                                processed_count += 1
                                logger.info(f"Successfully processed {file_path} into {len(chunks)} chunks")
                            else:
                                logger.warning(f"No chunks extracted from {file_path}")
                                error_count += 1
                        else:
                            logger.debug(f"Skipping unchanged file: {file_path}")
                            skipped_count += 1
                    except Exception as e:
                        logger.error(f"Error processing file {file_path}: {str(e)}\n{traceback.format_exc()}")
                        error_count += 1

            if files_processed:
                logger.info(f"Processing summary: {processed_count} files processed, {skipped_count} files skipped, {error_count} errors")
                logger.info(f"Total chunks created: {len(all_chunks)}")
                try:
                    logger.info("Creating vector store with embeddings...")
                    self.vectorstore = DocArrayInMemorySearch.from_texts(
                        texts=all_chunks,
                        embedding=self.embeddings
                    )
                    self._save_processed_files()
                    logger.info("Vector store created successfully")
                except Exception as e:
                    logger.error(f"Error creating vector store: {str(e)}\n{traceback.format_exc()}")
            else:
                logger.info(f"Scan complete: {skipped_count} files skipped, {error_count} errors, no new files to process")
        except Exception as e:
            logger.error(f"Error during document scan: {str(e)}\n{traceback.format_exc()}")

    def split_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        try:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )
            chunks = text_splitter.split_text(text)
            logger.debug(f"Split text into {len(chunks)} chunks")
            return chunks
        except Exception as e:
            logger.error(f"Error splitting text: {str(e)}\n{traceback.format_exc()}")
            return []

    def load_document(self, file_path: str) -> List[str]:
        """Load and extract text from a document."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            logger.info(f"Loading document: {file_path} (type: {file_extension})")
            if file_extension == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text += page.extract_text() + "\n"
                        logger.debug(f"Extracted text from page {page_num}")
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                logger.debug("Read text file successfully")
            elif file_extension in ['.doc', '.docx']:
                doc = docx.Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                logger.debug(f"Extracted text from {len(doc.paragraphs)} paragraphs")
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            chunks = self.split_text(text)
            logger.info(f"Successfully loaded {file_path} and created {len(chunks)} chunks")
            return chunks
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}\n{traceback.format_exc()}")
            return []

    def process_documents(self):
        """Process all documents and create embeddings."""
        self.scan_and_process_documents()

    def get_relevant_context(self, query: str, k: int = 3) -> List[Dict]:
        """Retrieve relevant context for a query."""
        try:
            logger.info(f"Getting relevant context for query: {query[:50]}...")
            if not self.vectorstore:
                logger.info("Vector store not initialized. Processing documents first...")
                self.process_documents()
                if not self.vectorstore:
                    logger.warning("Failed to initialize vector store")
                    return []

            # Search for similar documents
            docs = self.vectorstore.similarity_search(query, k=k)
            logger.info(f"Found {len(docs)} relevant documents")
            return [{"content": doc.page_content, "metadata": doc.metadata} for doc in docs]
        except Exception as e:
            logger.error(f"Error retrieving context: {str(e)}\n{traceback.format_exc()}")
            return [] 